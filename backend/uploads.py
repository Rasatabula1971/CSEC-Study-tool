# PHASE: build
"""
backend/uploads.py
==================
Upload-staging pipeline for the "Upload Material" feature (sessions 1-4).

This is build-phase content preparation -- the student-facing runtime never
touches it. A file dropped in the browser is *staged* (written to the SSD under
06_UPLOAD_STAGING and recorded in the upload_staging table) and its text is
extracted for preview. Nothing is ingested, classified, or embedded here.

Session 2 adds:
  * page-level OCR fallback in PDF extraction (PyMuPDF render -> Tesseract) for
    scanned pages -- both empty-string pages and "hidden scans" (a page that
    yields only a barcode / page number, below PAGE_TEXT_THRESHOLD chars);
  * full-file OCR when the whole file averages below FILE_AVG_THRESHOLD chars/page;
  * standalone image upload (.png/.jpg/.jpeg) via Tesseract directly;
  * chunked storage (upload_staging_chunks) for files past 500k chars, so the
    full content survives for session-3 classification instead of being truncated;
  * OCR quality signals (ocr_used / ocr_pages_count / ocr_confidence_avg /
    total_pages / truncated) recorded on the staging row.

Public surface:
  * stage_file(db, subject_id, original_name, file_bytes, file_type) -> staging_id
  * extract_text(staging_id, db) -> preview text (or None on failure)
  * reset_for_reextract(db, staging_id) -> None   (clears state for a re-run)
  * count_chunks(db, staging_id) -> int
  * get_staging_list(db, subject_id) -> list[dict]   (no full text -- list view)
  * get_staging_detail(db, staging_id) -> dict | None (full text included)

Extraction is deliberately its own, marker-rich implementation (page markers,
OCR markers, table markers) rather than reusing backend/extract.py, whose output
is plain joined text for the notes-classify flow. Tesseract is located via the
same drive-letter-agnostic resolver extract.py uses (TESSERACT_CMD env, then the
SSD-bundled binary, then PATH).
"""

import logging
import os
import re
from pathlib import Path

# Shared OCR primitives (binary resolution, DPI bomb-guard, page/image OCR) live in
# ocr_utils so uploads and the ingest_v2 adapters share ONE path. backend/ is on
# sys.path before uploads is imported (app.py / tests), and ocr_utils is light
# (heavy deps are lazy inside it), so this top-level import is safe.
from ocr_utils import OCR_DPI, ensure_tesseract, ocr_image, ocr_page, render_dpi

logger = logging.getLogger("csec.uploads")

MAX_EXTRACT_CHARS = 500_000        # preview cap stored in upload_staging.extracted_text
TRUNCATION_MARKER = "\n[Truncated: file exceeds 500k char limit]"  # legacy (session 1)
TRUNCATE_PREVIEW_MARKER = "\n[Truncated: see chunks]"             # appended to the preview
TRUNCATE_HARDCAP_MARKER = "\n[Truncated: 5M char hard cap]"

# These are intentionally HIGHER and richer than the ingest_v2 adapters' shared
# ocr_utils.OCR_TRIGGER_THRESHOLD (30): the student-upload path is purpose-tuned to OCR
# aggressively for phone-photos / scans, and pairs a per-page threshold with a
# whole-file-average mode. Deliberately NOT consolidated with the adapter trigger.
PAGE_TEXT_THRESHOLD = 50           # chars on a page below this -> OCR that page
FILE_AVG_THRESHOLD  = 100          # whole-file avg chars/page below this -> OCR every page
CHUNK_SIZE          = 100_000      # chars per row in upload_staging_chunks
MAX_TOTAL_CHARS     = 5_000_000    # hard cap on full extracted text, even with chunks
# OCR_LANG / OCR_DPI / PIL_PIXEL_LIMIT now live in ocr_utils (OCR_DPI re-imported above).

VALID_FILE_TYPES = ("pdf", "docx", "image")

# Everything outside this set is replaced with an underscore during sanitisation.
_BAD_CHARS = re.compile(r"[^a-zA-Z0-9._-]")

# Back-compat alias: the page-OCR primitive moved to ocr_utils.ocr_page. Retained so
# tests that call uploads._ocr_page directly keep resolving to the shared implementation.
_ocr_page = ocr_page


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _staging_root() -> Path:
    """Return {SSD_ROOT}/06_UPLOAD_STAGING. Raises IOError if SSD_ROOT is unset."""
    ssd_root = os.getenv("SSD_ROOT")
    if not ssd_root:
        raise IOError("SSD_ROOT is not set -- cannot locate the upload-staging area.")
    return Path(ssd_root) / "06_UPLOAD_STAGING"


def _subject_is_locked(db, subject_id: str) -> bool:
    row = db.execute(
        "SELECT 1 FROM subjects WHERE subject_id = ? AND syllabus_locked = 1",
        (subject_id,),
    ).fetchone()
    return row is not None


def safe_filename(original_name: str) -> str:
    """Produce a filesystem-safe filename that can never escape the staging dir.

    Rules:
      * path separators (/ and \\) become underscores
      * anything outside [a-zA-Z0-9._-] becomes an underscore
      * dots inside the stem are collapsed to underscores (kills '..' traversal),
        the single extension dot is preserved
      * extension is lower-cased and kept
      * capped at 100 chars, truncated FROM THE START so the extension survives
    """
    raw = (original_name or "").strip() or "file"
    raw = raw.replace("\\", "_").replace("/", "_")  # strip path separators

    ext = Path(raw).suffix.lower()                  # ".pdf" / ".docx" / "" / ...
    stem = raw[: len(raw) - len(ext)] if ext else raw

    stem = _BAD_CHARS.sub("_", stem)
    stem = stem.replace(".", "_")                   # no dots in the stem -> no '..'
    stem = stem.strip("._-") or "file"

    ext = _BAD_CHARS.sub("_", ext)                  # sanitise but keep the leading dot
    safe = stem + ext

    if len(safe) > 100:
        keep = 100 - len(ext)
        safe = (stem[-keep:] + ext) if keep > 0 else ext[-100:]
    return safe


# OCR plumbing (ensure_tesseract / ocr_image / render_dpi / ocr_page) moved to
# ocr_utils.py -- imported at the top of this module and shared with the ingest_v2
# adapters. The _extract_pdf / _extract_image extractors below call those directly.


def _finalize_extraction(full_text: str, total_pages, ocr_pages: list,
                         ocr_confidences: list) -> dict:
    """Apply the 5M hard cap, derive the truncated flag, slice chunks past 500k, and
    build the common extractor return dict. `text` is the <=500k preview that lands in
    upload_staging.extracted_text; `chunks` (only when truncated) holds the FULL text in
    CHUNK_SIZE slices for upload_staging_chunks."""
    if len(full_text) > MAX_TOTAL_CHARS:
        full_text = full_text[:MAX_TOTAL_CHARS] + TRUNCATE_HARDCAP_MARKER

    truncated = len(full_text) > MAX_EXTRACT_CHARS
    chunks = None
    if truncated:
        chunks = [full_text[i:i + CHUNK_SIZE]
                  for i in range(0, len(full_text), CHUNK_SIZE)]

    avg_conf = (int(sum(ocr_confidences) / len(ocr_confidences))
                if ocr_confidences else None)
    preview = full_text[:MAX_EXTRACT_CHARS] + (TRUNCATE_PREVIEW_MARKER if truncated else "")
    return {
        "text": preview,
        "total_pages": total_pages,
        "ocr_pages": ocr_pages,
        "ocr_confidence_avg": avg_conf,
        "truncated": truncated,
        "chunks": chunks,
    }


# ---------------------------------------------------------------------------
# Extractors  (each returns the common dict -- see _finalize_extraction)
# ---------------------------------------------------------------------------
def _extract_pdf(path) -> dict:
    """PyMuPDF extraction with OCR fallback. Native text is tried for every page; a
    page below PAGE_TEXT_THRESHOLD chars (or every page, when the whole file averages
    below FILE_AVG_THRESHOLD) is re-read via Tesseract and marked '[Page N - OCR]'."""
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    try:
        total_pages = doc.page_count
        native = []
        for pno in range(total_pages):
            page = doc.load_page(pno)
            native.append((pno + 1, page.get_text("text").strip(), page))

        total_native_chars = sum(len(t) for _, t, _ in native)
        avg_chars_per_page = total_native_chars / max(1, len(native))
        ocr_all = avg_chars_per_page < FILE_AVG_THRESHOLD

        parts, ocr_pages, ocr_confidences = [], [], []
        dpi_reduced = False
        for i, txt, page in native:
            needs_ocr = ocr_all or len(txt) < PAGE_TEXT_THRESHOLD
            if needs_ocr:
                if render_dpi(page)[1]:
                    dpi_reduced = True
                ocr_text, conf = ocr_page(page)
                parts.append(f"\n[Page {i} - OCR]\n{ocr_text}")
                ocr_pages.append(i)
                if conf is not None:
                    ocr_confidences.append(conf)
            elif not txt:
                parts.append(f"\n[Page {i} - no text]\n")
            else:
                parts.append(f"\n[Page {i}]\n{txt}")

        result = _finalize_extraction("".join(parts), total_pages, ocr_pages, ocr_confidences)
        result["ocr_dpi_reduced"] = dpi_reduced
        return result
    finally:
        doc.close()


def _extract_image(path) -> dict:
    """Standalone image OCR (.png/.jpg/.jpeg) via Tesseract."""
    ensure_tesseract()
    from PIL import Image
    img = Image.open(path)
    try:
        text, mean_conf = ocr_image(img)
    finally:
        img.close()
    return {
        "text": text,
        "total_pages": 1,
        "ocr_pages": [1],
        "ocr_confidence_avg": mean_conf,
        "truncated": False,
        "chunks": None,
    }


def _extract_docx(path) -> dict:
    """python-docx: paragraphs joined by blank lines, tables wrapped in markers."""
    import docx  # python-docx

    document = docx.Document(path)
    body = "\n\n".join(p.text for p in document.paragraphs)

    table_blocks = []
    for table in document.tables:
        rows = [" | ".join(cell.text for cell in row.cells) for row in table.rows]
        table_blocks.append("\n[Table]\n" + "\n".join(rows) + "\n[/Table]\n")

    if table_blocks:
        body = (body + "\n\n" if body else "") + "\n".join(table_blocks)

    # docx has no pages; OCR never applies. Truncation/chunking still does.
    return _finalize_extraction(body, None, [], [])


# ---------------------------------------------------------------------------
# Public surface
# ---------------------------------------------------------------------------
def stage_file(db, subject_id: str, original_name: str, file_bytes: bytes,
               file_type: str) -> int:
    """Validate, write to the SSD, and record one staged file. Returns staging_id.

    The file lands at
        {SSD_ROOT}/06_UPLOAD_STAGING/{subject_id}/{staging_id}_{safe_name}
    The row is inserted first (to mint the autoincrement staging_id), then the
    bytes are written, then stored_path is set -- all in one transaction, so a
    failed SSD write rolls the row back (no orphan row, no orphan file).

    Raises ValueError for an unlocked subject or an unsupported file_type, and
    IOError if the SSD write fails.
    """
    file_type = (file_type or "").lower()
    if file_type not in VALID_FILE_TYPES:
        raise ValueError(
            f"Unsupported file_type '{file_type}'. Expected one of: "
            f"{', '.join(VALID_FILE_TYPES)}."
        )
    if not _subject_is_locked(db, subject_id):
        raise ValueError(f"Subject '{subject_id}' is not a locked subject.")

    safe_name = safe_filename(original_name)
    subject_dir = _staging_root() / subject_id

    cur = db.execute(
        """
        INSERT INTO upload_staging
            (subject_id, original_name, stored_path, file_type, file_size_bytes,
             extract_status, status, updated_at)
        VALUES (?, ?, ?, ?, ?, 'pending', 'staged', datetime('now'))
        """,
        (subject_id, original_name, "", file_type, len(file_bytes)),
    )
    staging_id = cur.lastrowid
    stored_path = subject_dir / f"{staging_id}_{safe_name}"

    try:
        subject_dir.mkdir(parents=True, exist_ok=True)
        stored_path.write_bytes(file_bytes)
    except OSError as exc:
        db.rollback()  # discard the uncommitted row -- no orphan
        raise IOError(f"Failed to write staged file to {stored_path}: {exc}") from exc

    db.execute(
        "UPDATE upload_staging SET stored_path = ? WHERE staging_id = ?",
        (str(stored_path), staging_id),
    )
    db.commit()
    return staging_id


def extract_text(staging_id: int, db):
    """Extract text from a staged file, driving the extract_status state machine.

    pending -> extracting -> ready (extracted_text + OCR signals populated) | failed
    (extract_error populated, extracted_text left null). When the full text exceeds
    500k chars the file is flagged truncated and the FULL text is written to
    upload_staging_chunks in the same transaction as the staging-row UPDATE. Returns
    the preview text on success, or None if the row is missing or extraction fails.
    """
    row = db.execute(
        "SELECT stored_path, file_type FROM upload_staging WHERE staging_id = ?",
        (staging_id,),
    ).fetchone()
    if row is None:
        return None

    db.execute(
        "UPDATE upload_staging SET extract_status = 'extracting', "
        "updated_at = datetime('now') WHERE staging_id = ?",
        (staging_id,),
    )
    db.commit()

    try:
        # Dispatch via direct names (resolved through the module namespace at call
        # time) so the extractors stay monkeypatchable in tests.
        ft = row["file_type"]
        if ft == "pdf":
            result = _extract_pdf(row["stored_path"])
        elif ft == "docx":
            result = _extract_docx(row["stored_path"])
        elif ft == "image":
            result = _extract_image(row["stored_path"])
        else:
            raise ValueError(f"Unsupported file_type '{ft}'.")

        text = result["text"]
        ocr_pages = result.get("ocr_pages") or []
        ocr_pages_count = len(ocr_pages)
        ocr_used = 1 if ocr_pages_count > 0 else 0
        truncated = 1 if result.get("truncated") else 0
        ocr_dpi_reduced = 1 if result.get("ocr_dpi_reduced") else 0
        chunks = result.get("chunks")

        # A re-extract may have left stale chunks; clear them inside this transaction.
        db.execute("DELETE FROM upload_staging_chunks WHERE staging_id = ?", (staging_id,))
        db.execute(
            "UPDATE upload_staging SET extract_status = 'ready', extracted_text = ?, "
            "extract_error = NULL, ocr_used = ?, ocr_pages_count = ?, "
            "ocr_confidence_avg = ?, total_pages = ?, truncated = ?, "
            "ocr_dpi_reduced = ?, updated_at = datetime('now') WHERE staging_id = ?",
            (text, ocr_used, ocr_pages_count, result.get("ocr_confidence_avg"),
             result.get("total_pages"), truncated, ocr_dpi_reduced, staging_id),
        )
        if chunks:
            for idx, chunk in enumerate(chunks):
                db.execute(
                    "INSERT INTO upload_staging_chunks "
                    "(staging_id, chunk_index, chunk_text, ocr_used) VALUES (?, ?, ?, ?)",
                    (staging_id, idx, chunk, ocr_used),
                )
        db.commit()
        return text
    except Exception as exc:  # noqa: BLE001 -- recorded as a failed extraction
        db.rollback()
        db.execute(
            "UPDATE upload_staging SET extract_status = 'failed', "
            "extract_error = ?, extracted_text = NULL, "
            "updated_at = datetime('now') WHERE staging_id = ?",
            (str(exc)[:1000], staging_id),
        )
        db.commit()
        return None


def reset_for_reextract(db, staging_id: int) -> None:
    """Clear a staged row's extraction state so it can be re-run from scratch:
    status -> 'pending', preview + OCR signals cleared, any chunks deleted."""
    db.execute("DELETE FROM upload_staging_chunks WHERE staging_id = ?", (staging_id,))
    db.execute(
        "UPDATE upload_staging SET extract_status = 'pending', extracted_text = NULL, "
        "extract_error = NULL, ocr_used = 0, ocr_pages_count = 0, "
        "ocr_confidence_avg = NULL, truncated = 0, ocr_dpi_reduced = 0, "
        "updated_at = datetime('now') WHERE staging_id = ?",
        (staging_id,),
    )
    db.commit()


def count_chunks(db, staging_id: int) -> int:
    """Number of upload_staging_chunks rows held for a staged file."""
    return db.execute(
        "SELECT COUNT(*) FROM upload_staging_chunks WHERE staging_id = ?",
        (staging_id,),
    ).fetchone()[0]


def get_staging_list(db, subject_id: str) -> list:
    """List view for one subject, newest first. Excludes the full extracted text
    (too large for a list) -- only its length, plus the badge signals the UI needs
    (ocr_used / ocr_confidence_avg / truncated)."""
    rows = db.execute(
        """
        SELECT staging_id, original_name, file_type, file_size_bytes,
               extract_status, status, created_at,
               ocr_used, ocr_confidence_avg, truncated, ocr_dpi_reduced,
               CASE WHEN extracted_text IS NULL THEN NULL
                    ELSE length(extracted_text) END AS extracted_text_length
        FROM   upload_staging
        WHERE  subject_id = ?
        ORDER  BY created_at DESC, staging_id DESC
        """,
        (subject_id,),
    ).fetchall()
    out = []
    for r in rows:
        d = dict(r)
        d["ocr_used"] = bool(d.get("ocr_used"))
        d["truncated"] = bool(d.get("truncated"))
        d["ocr_dpi_reduced"] = bool(d.get("ocr_dpi_reduced"))
        out.append(d)
    return out


def get_staging_detail(db, staging_id: int):
    """Full row for one staged file, INCLUDING extracted_text and OCR signals.
    None if absent."""
    row = db.execute(
        """
        SELECT staging_id, subject_id, original_name, file_type, file_size_bytes,
               stored_path, extract_status, extract_error, extracted_text,
               status, created_at,
               ocr_used, ocr_pages_count, ocr_confidence_avg, total_pages, truncated,
               ocr_dpi_reduced
        FROM   upload_staging
        WHERE  staging_id = ?
        """,
        (staging_id,),
    ).fetchone()
    return dict(row) if row else None
