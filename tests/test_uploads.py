"""
tests/test_uploads.py
=====================
Upload session 1 -- unit tests for backend/uploads.py.

Real in-memory SQLite (schema.sql + apply_runtime_migrations + sqlite-vec, FKs
ON) so the upload_staging CHECK constraints and FKs are genuinely exercised.
SSD_ROOT is pointed at a per-test tempdir so staged files land somewhere
disposable -- never the real SSD. No Ollama.

  1. stage_file writes a row and a file on disk.
  2. stage_file rejects an unlocked subject (no row, no file).
  3. stage_file rejects a bad file_type.
  4. extract_text(PDF) -> ready + text + page marker.
  5. extract_text(DOCX) -> ready + paragraph text + table markers.
  6. extract_text on a corrupt file -> failed + error, text null.
  7. Filename sanitisation strips traversal/path chars, keeps the extension.
  8. 500k-char cap with a visible truncation marker.

Run: pytest tests/test_uploads.py -v
"""

import io
import os
import sqlite3
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "backend"))

SCHEMA_PATH = ROOT / "backend" / "db" / "schema.sql"

import app as app_module   # noqa: E402  (apply_runtime_migrations)
import uploads             # noqa: E402

SUBJECT = "Principles_of_Business"


# --- fixtures --------------------------------------------------------------
def open_test_db() -> sqlite3.Connection:
    try:
        import sqlite_vec
    except ImportError:
        pytest.skip("sqlite-vec not installed -- skipping upload tests")
    db = sqlite3.connect(":memory:", check_same_thread=False)
    db.enable_load_extension(True)
    sqlite_vec.load(db)
    db.enable_load_extension(False)
    db.execute("PRAGMA foreign_keys = ON")
    db.row_factory = sqlite3.Row
    for stmt in SCHEMA_PATH.read_text(encoding="utf-8").split(";"):
        if stmt.strip():
            db.execute(stmt)
    db.commit()
    app_module.apply_runtime_migrations(db)  # creates upload_staging (m012)
    return db


def seed(db: sqlite3.Connection, locked: int = 1) -> None:
    db.execute(
        "INSERT INTO subjects (subject_id, display_name, syllabus_locked) VALUES (?, ?, ?)",
        (SUBJECT, "Principles of Business", locked),
    )
    db.commit()


@pytest.fixture
def db(tmp_path, monkeypatch):
    # Point the staging area at a disposable tempdir BEFORE migrations run
    # (apply_runtime_migrations -> ensure_staging_dirs reads SSD_ROOT).
    monkeypatch.setenv("SSD_ROOT", str(tmp_path))
    conn = open_test_db()
    seed(conn)
    yield conn
    conn.close()


# --- builders --------------------------------------------------------------
def make_pdf_bytes(text: str = "Hello PDF extraction test") -> bytes:
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    data = doc.tobytes()
    doc.close()
    return data


def make_docx_bytes(paragraphs=("First paragraph alpha", "Second paragraph beta"),
                    table=(("A1", "A2"), ("B1", "B2"))) -> bytes:
    import docx
    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    if table:
        t = document.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, val in enumerate(row):
                t.cell(r, c).text = val
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# --- tests -----------------------------------------------------------------
def test_stage_file_writes_row_and_file(db):
    sid = uploads.stage_file(db, SUBJECT, "test.pdf", make_pdf_bytes(), "pdf")
    assert isinstance(sid, int)

    row = db.execute(
        "SELECT * FROM upload_staging WHERE staging_id = ?", (sid,)
    ).fetchone()
    assert row["extract_status"] == "pending"
    assert row["status"] == "staged"
    assert row["file_type"] == "pdf"
    assert row["original_name"] == "test.pdf"
    assert os.path.exists(row["stored_path"]), "file written to the SSD staging area"


def test_stage_file_rejects_unlocked_subject(db):
    db.execute("UPDATE subjects SET syllabus_locked = 0 WHERE subject_id = ?", (SUBJECT,))
    db.commit()
    with pytest.raises(ValueError):
        uploads.stage_file(db, SUBJECT, "test.pdf", make_pdf_bytes(), "pdf")
    assert db.execute("SELECT COUNT(*) FROM upload_staging").fetchone()[0] == 0


def test_stage_file_rejects_bad_file_type(db):
    with pytest.raises(ValueError):
        uploads.stage_file(db, SUBJECT, "sheet.xlsx", b"junk", "xlsx")
    assert db.execute("SELECT COUNT(*) FROM upload_staging").fetchone()[0] == 0


def test_extract_pdf_populates_text_and_status(db):
    sid = uploads.stage_file(db, SUBJECT, "test.pdf",
                             make_pdf_bytes("Hello PDF extraction test"), "pdf")
    out = uploads.extract_text(sid, db)

    row = db.execute(
        "SELECT extract_status, extracted_text FROM upload_staging WHERE staging_id = ?",
        (sid,),
    ).fetchone()
    assert row["extract_status"] == "ready"
    assert "Hello PDF extraction test" in row["extracted_text"]
    assert "[Page 1]" in row["extracted_text"]
    assert out == row["extracted_text"]


def test_extract_docx_populates_text_and_markers(db):
    sid = uploads.stage_file(db, SUBJECT, "notes.docx", make_docx_bytes(), "docx")
    uploads.extract_text(sid, db)

    row = db.execute(
        "SELECT extract_status, extracted_text FROM upload_staging WHERE staging_id = ?",
        (sid,),
    ).fetchone()
    assert row["extract_status"] == "ready"
    text = row["extracted_text"]
    assert "First paragraph alpha" in text
    assert "Second paragraph beta" in text
    assert "[Table]" in text and "[/Table]" in text
    assert "A1 | A2" in text  # row cells joined with " | "


def test_extract_corrupt_file_sets_failed(db):
    sid = uploads.stage_file(db, SUBJECT, "broken.pdf", b"not a real pdf at all", "pdf")
    out = uploads.extract_text(sid, db)
    assert out is None

    row = db.execute(
        "SELECT extract_status, extract_error, extracted_text FROM upload_staging "
        "WHERE staging_id = ?", (sid,),
    ).fetchone()
    assert row["extract_status"] == "failed"
    assert row["extract_error"]                 # populated
    assert row["extracted_text"] is None


def test_filename_sanitisation_strips_traversal(db):
    sid = uploads.stage_file(db, SUBJECT, "../../../etc/passwd.pdf", make_pdf_bytes(), "pdf")
    row = db.execute(
        "SELECT stored_path FROM upload_staging WHERE staging_id = ?", (sid,)
    ).fetchone()
    stored_path = row["stored_path"]
    filename = os.path.basename(stored_path)

    assert ".." not in stored_path
    assert "/" not in filename
    assert "\\" not in filename
    assert filename.endswith(".pdf")


def test_500k_char_cap_on_extraction(db):
    big = "x" * 600_000
    sid = uploads.stage_file(db, SUBJECT, "huge.docx",
                             make_docx_bytes(paragraphs=(big,), table=None), "docx")
    uploads.extract_text(sid, db)

    row = db.execute(
        "SELECT extracted_text FROM upload_staging WHERE staging_id = ?", (sid,)
    ).fetchone()
    text = row["extracted_text"]
    assert uploads.TRUNCATION_MARKER in text
    assert len(text) == uploads.MAX_EXTRACT_CHARS + len(uploads.TRUNCATION_MARKER)
